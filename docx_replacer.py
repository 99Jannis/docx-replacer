from docx import Document
from os import path

# file path of input and output document
input_file_path = 'test.docx'
output_file_path = 'result_'+input_file_path 

# searchterms and replacements, replace it with your own terms
#   Note: searchterm and replacement are case-sensitive(!)
#   keys (left) are searchterm and values (right) are replacements
#   Unicodes are in format \uXXXX
variables = {
    "searchterm" : "replacement",
    "\u201E" : "\u00BB", # double low quotemark : guillemot right
    "\u201C" : "\u00AB", # left double quotemark : guillemot left
}


def main():
    # check if file is valid and can be opened
    if is_docx is False:
        return -1
    try: 
        document = Document(input_file_path)
    except:
        print("ERROR: Cannot open file")
        return -1
    
    # replace text in document
    for searchterm, replacement in variables.items():
        # replace text in paragraphs
        for paragraph in document.paragraphs:
            replace_text_in_paragraph(paragraph, searchterm, replacement)

        # replace text in tables
        for table in document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, searchterm, replacement)
    # save results as new document
    document.save(output_file_path)


def replace_text_in_paragraph(paragraph, searchterm, replacement)->None:
    """
    replaces text in a given paragraph
    """
    if searchterm in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if searchterm in item.text:
                item.text = item.text.replace(searchterm, replacement)
                
def is_docx(filepath:str)->bool:
    """
    checks if file is docx format and if path is valid

    Args:
        filepath (str): path of file

    Returns:
        bool: True, if everthing is valid. False, otherwise
    """
    if not path.isdir(filepath):
        print("ERROR: Path does not exist!")
        return False
    if not filepath.endswith(".docx"):
        print("ERROR: File is not word")
        return False
    return True


if __name__ == '__main__':
    main()
